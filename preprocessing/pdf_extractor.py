"""
pdf_extractor.py
-----------------
Extract raw text from resume files (PDF or DOCX) so they can be fed
into the cleaning and embedding pipeline.

Usage:
    from preprocessing.pdf_extractor import extract_text

    text = extract_text("resume.pdf")
    text = extract_text("resume.docx")
"""

import os
import pdfplumber
import docx


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file, page by page."""
    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file, paragraph by paragraph."""
    document = docx.Document(file_path)
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_text(file_path: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from in-memory file bytes (useful for FastAPI UploadFile,
    where the file arrives as bytes rather than a saved path).
    """
    import io

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text_chunks = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
        return "\n".join(text_chunks)

    elif ext == ".docx":
        document = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in document.paragraphs if p.text.strip())

    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


if __name__ == "__main__":
    # Quick manual test — point this at a real file in data/raw/ to verify
    sample_path = "data/raw/resume_dataset/data/data/HR/10399912.pdf"
    if os.path.exists(sample_path):
        print(extract_text(sample_path)[:500])
    else:
        print(f"Sample file not found at {sample_path}. Update the path to test.")